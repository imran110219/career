"""Build the designed Senior Backend Engineer resume without changing canonical sources."""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "resumes" / "senior-backend"
RESUME_SOURCE = OUTPUT_DIR / "resume.md"
DOCX_OUT = OUTPUT_DIR / "resume.docx"
PDF_OUT = OUTPUT_DIR / "resume.pdf"

NAVY = RGBColor(27, 55, 92)
CHARCOAL = RGBColor(35, 39, 45)
MUTED = RGBColor(82, 88, 96)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_hyperlink(paragraph, text: str, url: str, *, color="1B375C"):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    properties.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_bottom_border(paragraph, color="1B375C"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_keep(paragraph, *, with_next=False, together=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if with_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if together:
        p_pr.append(OxmlElement("w:keepLines"))


def add_run(paragraph, text, *, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def configure_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.03

    for name, size, color in (("Resume Section", 11.3, NAVY), ("Resume Role", 10.3, CHARCOAL)):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    return document


def section_heading(document: Document, text: str):
    paragraph = document.add_paragraph(style="Resume Section")
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text.upper())
    set_bottom_border(paragraph)
    set_keep(paragraph, with_next=True)
    return paragraph


def bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.style = document.styles["Normal"]
    paragraph.paragraph_format.left_indent = Inches(0.17)
    paragraph.paragraph_format.first_line_indent = Inches(-0.13)
    paragraph.paragraph_format.space_after = Pt(1.8)
    paragraph.paragraph_format.line_spacing = 1.02
    add_run(paragraph, "•  ", color=NAVY)
    add_run(paragraph, text)
    set_keep(paragraph, together=True)
    return paragraph


def add_experience(document: Document, company: str, location: str, role: str, dates: str, bullets: list[str]):
    heading = document.add_paragraph(style="Resume Role")
    heading.paragraph_format.space_before = Pt(4.5)
    heading.paragraph_format.space_after = Pt(0.4)
    add_run(heading, company, bold=True)
    add_run(heading, f"  |  {location}", color=MUTED)
    set_keep(heading, with_next=True)

    details = document.add_paragraph()
    details.paragraph_format.space_after = Pt(1.5)
    add_run(details, role, bold=True, size=9.7)
    add_run(details, f"  |  {dates}", size=9.7, color=MUTED)
    set_keep(details, with_next=True)

    for item in bullets:
        bullet(document, item)


def load_resume(source: Path) -> dict:
    """Read the tailored Markdown source used by the finalized Senior Backend layout."""
    lines = source.read_text().splitlines()
    title_index = next(index for index, line in enumerate(lines) if line.startswith("# "))
    name = lines[title_index][2:].strip()
    first_section = next(index for index, line in enumerate(lines) if line.startswith("## "))
    header = [line.strip() for line in lines[title_index + 1:first_section] if line.strip()]
    target_title, contact_line, link_line = header

    sections: dict[str, list[str]] = {}
    active = None
    for line in lines[first_section:]:
        if line.startswith("## "):
            active = line[3:].strip()
            sections[active] = []
        elif active is not None:
            sections[active].append(line.rstrip())

    def bullets(section: str) -> list[str]:
        return [line[2:].strip() for line in sections[section] if line.startswith("- ")]

    skills = []
    for item in bullets("Technical Skills"):
        match = re.fullmatch(r"\*\*(.+?):\*\*\s*(.+)", item)
        if not match:
            raise ValueError(f"Invalid skill line in {source}: {item}")
        skills.append(match.groups())

    experience = []
    current = None
    for line in sections["Professional Experience"]:
        if line.startswith("### "):
            company, role = line[4:].split(" - ", 1)
            current = {"company": company, "role": role, "location": "", "dates": "", "bullets": []}
            experience.append(current)
        elif current and line.strip() and not current["location"]:
            current["location"], current["dates"] = (part.strip() for part in line.split("|", 1))
        elif current and line.startswith("- "):
            current["bullets"].append(line[2:].strip())

    summary = " ".join(line.strip() for line in sections["Professional Summary"] if line.strip())
    education = next(line.strip() for line in sections["Education"] if line.strip())
    degree, institution_and_dates = education.split(", ", 1)
    institution, dates = (part.strip() for part in institution_and_dates.split("|", 1))
    location, phone, email = (part.strip() for part in contact_line.split("|"))
    links = {label.strip(): value.strip() for label, value in (part.strip().split(": ", 1) for part in link_line.split("|"))}
    return {
        "name": name,
        "target_title": target_title,
        "location": location,
        "phone": phone,
        "email": email,
        "linkedin": links["LinkedIn"],
        "website": links["Website"],
        "github": links["GitHub"],
        "summary": summary,
        "highlights": bullets("Career Highlights"),
        "skills": skills,
        "experience": experience,
        "degree": degree,
        "institution": institution,
        "education_dates": dates,
        "publications": bullets("Publications"),
    }


def build_docx(output: Path, source: Path = RESUME_SOURCE):
    content = load_resume(source)
    document = configure_document()

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    add_run(name, content["name"], bold=True, size=22, color=NAVY)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    add_run(title, content["target_title"].replace(" | ", "  |  "), bold=True, size=11.4, color=NAVY)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(0)
    add_run(contact, f"{content['location']}  |  {content['phone']}  |  ", size=9.4, color=MUTED)
    add_hyperlink(contact, content["email"], f"mailto:{content['email']}")

    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(3)
    add_hyperlink(links, content["linkedin"].removeprefix("www."), f"https://{content['linkedin']}")
    add_run(links, "  |  ", size=9.4, color=MUTED)
    add_hyperlink(links, content["website"].removeprefix("www."), f"https://{content['website']}")
    add_run(links, "  |  ", size=9.4, color=MUTED)
    add_hyperlink(links, content["github"].removeprefix("www."), f"https://{content['github']}")
    set_bottom_border(links)

    section_heading(document, "Professional Summary")
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(1)
    summary.paragraph_format.line_spacing = 1.03
    add_run(summary, content["summary"])

    section_heading(document, "Career Highlights")
    for item in content["highlights"]:
        bullet(document, item)

    section_heading(document, "Technical Skills")
    for label, value in content["skills"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        add_run(paragraph, f"{label}: ", bold=True, size=9.5, color=NAVY)
        add_run(paragraph, value, size=9.5)
        set_keep(paragraph, together=True)

    document.add_page_break()
    section_heading(document, "Professional Experience")
    for role in content["experience"]:
        add_experience(document, role["company"], role["location"], role["role"], role["dates"].replace(" - ", " – "), role["bullets"])

    section_heading(document, "Education")
    education = document.add_paragraph()
    education.paragraph_format.space_after = Pt(0)
    add_run(education, content["degree"], bold=True, size=9.8)
    add_run(education, f"  |  {content['institution']}  |  {content['education_dates'].replace(' - ', ' – ')}", size=9.8, color=MUTED)
    set_keep(education, together=True)

    section_heading(document, "Publications")
    for publication in content["publications"]:
        publication_paragraph = document.add_paragraph()
        publication_paragraph.paragraph_format.space_after = Pt(1)
        publication_paragraph.paragraph_format.line_spacing = 1.0
        add_run(publication_paragraph, publication, size=9.2)
        set_keep(publication_paragraph, together=True)

    document.save(output)


def export_pdf(docx: Path, pdf: Path):
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise RuntimeError("LibreOffice is required to export the PDF.")
    temporary = OUTPUT_DIR / ".pdf-export"
    temporary.mkdir(exist_ok=True)
    profile = Path("/tmp/career-resume-libreoffice-profile")
    profile.mkdir(exist_ok=True)
    subprocess.run([
        libreoffice,
        f"-env:UserInstallation={profile.as_uri()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(temporary),
        str(docx),
    ], check=True)
    generated = temporary / f"{docx.stem}.pdf"
    generated.replace(pdf)
    temporary.rmdir()


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx(DOCX_OUT)
    export_pdf(DOCX_OUT, PDF_OUT)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()

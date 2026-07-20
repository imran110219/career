import argparse
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
RESUME_MD = ROOT / "master" / "resume.md"
RTF_OUT = Path("/tmp/master_resume.rtf")
PDF_OUT = ROOT / "master" / "resume.pdf"


def rtf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")


def display_text(value: str) -> str:
    return value.replace("**", "").replace("`", "")


def parse_markdown(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            blocks.append(("title", stripped[2:].strip()))
        elif stripped.startswith("## "):
            blocks.append(("h1", stripped[3:].strip()))
        elif stripped.startswith("### "):
            blocks.append(("h2", stripped[4:].strip()))
        elif stripped.startswith("- [ ]"):
            blocks.append(("check", stripped[5:].strip()))
        elif stripped.startswith("- "):
            blocks.append(("bullet", stripped[2:].strip()))
        elif stripped.startswith("| "):
            body = stripped.replace("|", "").replace(" ", "")
            if body and set(body) == {"-"}:
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            blocks.append(("table", "  |  ".join(cells)))
        elif stripped.startswith("```"):
            continue
        else:
            blocks.append(("body", stripped))

    return blocks


def build_rtf(blocks: list[tuple[str, str]]) -> str:
    rtf = [
        r"{\rtf1\ansi\deff0{\fonttbl{\f0 Calibri;}}",
        r"{\colortbl ;\red46\green116\blue181;\red31\green77\blue120;\red85\green85\blue85;}",
        r"\paperw12240\paperh15840\margl1080\margr1080\margt900\margb900",
        r"\f0",
    ]

    for kind, raw_text in blocks:
        text = rtf_escape(display_text(raw_text))
        if kind == "title":
            rtf.append(r"\pard\qc\b\fs34\cf2 " + text + r"\b0\cf0\par")
        elif kind == "h1":
            rtf.append(r"\pard\sb220\sa90\b\fs24\cf1 " + text + r"\b0\cf0\par")
        elif kind == "h2":
            rtf.append(r"\pard\sb140\sa50\b\fs21\cf2 " + text + r"\b0\cf0\par")
        elif kind == "bullet":
            rtf.append(r"\pard\li360\fi-180\sa45\fs20 \'95\tab " + text + r"\par")
        elif kind == "check":
            rtf.append(r"\pard\li360\fi-180\sa45\fs20 [ ]\tab " + text + r"\par")
        elif kind == "table":
            rtf.append(r"\pard\sa45\fs19\cf3 " + text + r"\cf0\par")
        else:
            rtf.append(r"\pard\sa70\fs20 " + text + r"\par")

    rtf.append("}")
    return "\n".join(rtf)


def pdf_escape(value: str) -> str:
    return (
        value.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .encode("latin-1", "replace")
        .decode("latin-1")
    )


def build_pdf(blocks: list[tuple[str, str]], output: Path) -> None:
    page_width = 612
    page_height = 792
    margin_left = 54
    y_start = 744
    y_min = 54
    pages: list[list[str]] = []
    current: list[str] = []
    y = y_start

    def add_line(text: str, size: int = 10, leading: int = 13) -> None:
        nonlocal current, y
        if y < y_min:
            pages.append(current)
            current = []
            y = y_start
        safe = pdf_escape(text)
        current.append(f"BT /F1 {size} Tf {margin_left} {y} Td ({safe}) Tj ET")
        y -= leading

    for kind, text in blocks:
        text = display_text(text)
        if kind == "title":
            add_line(text, 18, 24)
        elif kind == "h1":
            y -= 4
            add_line(text.upper(), 12, 17)
        elif kind == "h2":
            add_line(text, 11, 15)
        elif kind == "bullet":
            for index, line in enumerate(wrap(text, width=92)):
                prefix = "- " if index == 0 else "  "
                add_line(prefix + line, 9, 12)
        elif kind == "check":
            for index, line in enumerate(wrap(text, width=90)):
                prefix = "[ ] " if index == 0 else "    "
                add_line(prefix + line, 9, 12)
        elif kind == "table":
            for line in wrap(text, width=96):
                add_line(line, 8, 11)
        else:
            for line in wrap(text, width=94):
                add_line(line, 9, 12)

    if current:
        pages.append(current)

    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    kids = " ".join(f"{3 + i * 2} 0 R" for i in range(len(pages)))
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())

    for page_index, page_lines in enumerate(pages):
        page_obj = 3 + page_index * 2
        content_obj = page_obj + 1
        stream = "\n".join(page_lines).encode("latin-1")
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 {3 + len(pages) * 2} 0 R >> >> "
            f"/Contents {content_obj} 0 R >>".encode()
        )
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )

    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode())
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    pdf.extend(b"0000000000 65535 f\n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n\n".encode())
    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    output.write_bytes(pdf)


def build_docx(blocks: list[tuple[str, str]], output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)

    for kind, raw_text in blocks:
        value = display_text(raw_text)
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            run.bold = True
            run.font.size = Pt(16)
        elif kind == "h1":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value.upper())
            run.bold = True
            run.font.size = Pt(11)
        elif kind == "h2":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(value)
            run.bold = True
            run.font.size = Pt(10)
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.add_run(value)
        else:
            document.add_paragraph(value)

    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a resume PDF and intermediate RTF from Markdown.")
    parser.add_argument("input", nargs="?", type=Path, default=RESUME_MD, help="Markdown resume source")
    parser.add_argument("--pdf-output", type=Path, default=PDF_OUT, help="Destination PDF path")
    parser.add_argument("--rtf-output", type=Path, default=RTF_OUT, help="Destination intermediate RTF path")
    parser.add_argument("--docx-output", type=Path, help="Optional destination DOCX path")
    args = parser.parse_args()

    blocks = parse_markdown(args.input.read_text())
    args.rtf_output.write_text(build_rtf(blocks))
    build_pdf(blocks, args.pdf_output)
    if args.docx_output:
        build_docx(blocks, args.docx_output)
    print(args.rtf_output)
    print(args.pdf_output)
    if args.docx_output:
        print(args.docx_output)


if __name__ == "__main__":
    main()

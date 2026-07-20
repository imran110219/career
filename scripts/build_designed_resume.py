"""Reusable, ATS-safe Markdown-to-DOCX/PDF resume builder using the approved design system."""

from __future__ import annotations

import argparse
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(27, 55, 92)
CHARCOAL = RGBColor(35, 39, 45)
MUTED = RGBColor(82, 88, 96)


def clean(value: str) -> str:
    return value.replace("**", "").replace("`", "").strip()


def blocks(markdown: str):
    for raw in markdown.splitlines():
        value = raw.strip()
        if not value or value.startswith("```"):
            continue
        if value.startswith("# "):
            yield "title", clean(value[2:])
        elif value.startswith("## "):
            yield "section", clean(value[3:])
        elif value.startswith("### "):
            yield "subsection", clean(value[4:])
        elif value.startswith("- [ ]"):
            continue
        elif value.startswith("- "):
            yield "bullet", clean(value[2:])
        elif value.startswith("|"):
            cells = [clean(cell) for cell in value.strip("|").split("|")]
            if all(not cell or set(cell) <= {"-", ":"} for cell in cells):
                continue
            yield "table", "  |  ".join(cells)
        else:
            yield "body", clean(value)


def font(run, size=None, bold=None, color=None):
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def keep(paragraph, next_item=False, lines=False):
    properties = paragraph._p.get_or_add_pPr()
    if next_item:
        properties.append(OxmlElement("w:keepNext"))
    if lines:
        properties.append(OxmlElement("w:keepLines"))


def divider(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "1B375C")
    borders.append(bottom)
    properties.append(borders)


def configure() -> Document:
    document = Document()
    page = document.sections[0]
    page.page_width, page.page_height = Inches(8.2677), Inches(11.6929)
    page.top_margin = page.bottom_margin = Inches(0.58)
    page.left_margin = page.right_margin = Inches(0.64)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.03
    for name, size, color in (("Designed Section", 11.3, NAVY), ("Designed Subsection", 10.4, CHARCOAL)):
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    return document


def add_header(document: Document, name: str, metadata: list[str]):
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(1)
    run = heading.add_run(name)
    font(run, 22, True, NAVY)
    for index, value in enumerate(metadata):
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(1 if index == len(metadata) - 1 else 0)
        run = line.add_run(value)
        font(run, 10.2 if index == 0 else 9.4, index == 0, NAVY if index == 0 else MUTED)
        if index == len(metadata) - 1:
            divider(line)


def build(input_file: Path, docx_file: Path, *, break_before_experience: bool = False):
    items = list(blocks(input_file.read_text()))
    document = configure()
    name = next((value for kind, value in items if kind == "title"), "Resume")
    first_section = next((index for index, (kind, _) in enumerate(items) if kind == "section"), len(items))
    metadata = [value for kind, value in items[1:first_section] if kind == "body"]
    add_header(document, name, metadata)
    start = first_section
    for kind, value in items[start:]:
        if break_before_experience and kind == "section" and value.casefold() == "professional experience":
            document.add_page_break()
        if kind == "section" and value.casefold() == "review checklist":
            break
        if kind == "section":
            paragraph = document.add_paragraph(style="Designed Section")
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.add_run(value.upper())
            divider(paragraph)
            keep(paragraph, next_item=True)
        elif kind == "subsection":
            paragraph = document.add_paragraph(style="Designed Subsection")
            paragraph.paragraph_format.space_before = Pt(4.5)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.add_run(value)
            keep(paragraph, next_item=True)
        elif kind == "bullet":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.17)
            paragraph.paragraph_format.first_line_indent = Inches(-0.13)
            paragraph.paragraph_format.space_after = Pt(1.8)
            paragraph.paragraph_format.line_spacing = 1.02
            marker = paragraph.add_run("•  ")
            font(marker, color=NAVY)
            paragraph.add_run(value)
            keep(paragraph, lines=True)
        elif kind == "table":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1.3)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            font(run, 9.3, False, MUTED)
            keep(paragraph, lines=True)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(value)
            keep(paragraph, lines=True)
    docx_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_file)


def export_pdf(docx_file: Path, pdf_file: Path):
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if not office:
        raise RuntimeError("LibreOffice is required for PDF export.")
    export_dir = pdf_file.parent / ".pdf-export"
    export_dir.mkdir(exist_ok=True)
    profile = Path("/tmp/career-resume-libreoffice-profile")
    profile.mkdir(exist_ok=True)
    subprocess.run([office, f"-env:UserInstallation={profile.as_uri()}", "--headless", "--convert-to", "pdf", "--outdir", str(export_dir), str(docx_file)], check=True)
    (export_dir / f"{docx_file.stem}.pdf").replace(pdf_file)
    export_dir.rmdir()


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("docx", type=Path)
    parser.add_argument("pdf", type=Path)
    parser.add_argument("--break-before-experience", action="store_true", help="Start Professional Experience on page two for a balanced two-page tailored resume.")
    args = parser.parse_args()
    build(args.input, args.docx, break_before_experience=args.break_before_experience)
    export_pdf(args.docx, args.pdf)


if __name__ == "__main__":
    main()
